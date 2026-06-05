"""DOCX extractor with python-docx — paragraphs + tables (optional dep)."""

from __future__ import annotations

import io

from docswarm.ingest.extractors.base import ExtractResult


def _table_to_markdown(table) -> tuple[str, list[list[str]]]:
    rows = [[(cell.text or "").strip() for cell in row.cells] for row in table.rows]
    if not rows:
        return "", []
    header, *body = rows
    md = ["| " + " | ".join(header) + " |",
          "| " + " | ".join("---" for _ in header) + " |"]
    for r in body:
        md.append("| " + " | ".join(r) + " |")
    return "\n".join(md), rows


def extract_docx(content: bytes) -> ExtractResult:
    try:
        import docx  # python-docx
    except ImportError:
        return ExtractResult(
            text="", method="unsupported",
            error="python-docx not installed. Run: pip install docswarm[ingest]")

    doc = docx.Document(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if (p.text or "").strip()]
    text = "\n\n".join(paragraphs)

    tables: list[dict] = []
    for t in doc.tables:
        md, rows = _table_to_markdown(t)
        if md:
            tables.append({"markdown": md, "rows": rows})

    return ExtractResult(
        text=text, method="docx",
        tables_markdown=[t["markdown"] for t in tables], tables=tables,
        chars=len(text), meta={"n_tables": len(tables),
                               "n_paragraphs": len(paragraphs)},
    )
